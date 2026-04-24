"""Generate the 25 Sprint O fixture .eml files from manifest.yaml.

Idempotent: running twice overwrites existing .eml files with byte-identical
output for the same manifest. Attachment payloads (invoice PDFs, contract
DOCXs, generic PDFs) are generated inline with reportlab / python-docx so
no binary blobs are tracked in git alongside the .eml files.

Run from repo root:

    .venv/Scripts/python.exe data/fixtures/emails_sprint_o/generate_fixtures.py
"""

from __future__ import annotations

import io
from email.message import EmailMessage
from pathlib import Path

import yaml
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

FIXTURE_DIR = Path(__file__).parent


def _invoice_pdf(invoice_number: int, amount_huf: int, supplier: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 780, "INVOICE / SZAMLA")
    c.setFont("Helvetica", 11)
    c.drawString(72, 750, f"Invoice No: INV-2026-{invoice_number:04d}")
    c.drawString(72, 734, f"Szamlaszam: INV-2026-{invoice_number:04d}")
    c.drawString(72, 718, f"Supplier / Kibocsato: {supplier}")
    c.drawString(72, 702, "Customer: Acme Kft, 1051 Budapest, Szent Istvan ter 1")
    c.drawString(72, 686, "Issue date: 2026-04-01  |  Due date: 2026-05-15")
    c.drawString(72, 660, "Description")
    c.drawString(400, 660, "Amount")
    c.line(72, 655, 520, 655)
    c.drawString(72, 640, "Consulting services, March 2026")
    c.drawString(400, 640, f"{amount_huf:,} HUF")
    c.drawString(72, 620, "VAT 27%")
    c.drawString(400, 620, f"{int(amount_huf * 0.27):,} HUF")
    c.line(72, 610, 520, 610)
    c.drawString(72, 592, "Total / Osszesen")
    c.drawString(400, 592, f"{int(amount_huf * 1.27):,} HUF")
    c.drawString(72, 560, "Payment: wire transfer, IBAN HU12 1111 2222 3333 4444 5555 6666")
    c.drawString(72, 544, "Please reference the invoice number on the transfer.")
    c.showPage()
    c.save()
    return buf.getvalue()


def _generic_pdf(title: str, body_lines: list[str]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 780, title)
    c.setFont("Helvetica", 10)
    y = 750
    for line in body_lines:
        c.drawString(72, y, line[:110])
        y -= 14
        if y < 80:
            break
    c.showPage()
    c.save()
    return buf.getvalue()


def _contract_docx(title: str, clauses: list[str]) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(
        "This service agreement (the 'Agreement') is entered into "
        "between Acme Kft and the undersigned Contractor, effective "
        "as of 2026-04-01."
    )
    for i, clause in enumerate(clauses, 1):
        doc.add_heading(f"Section {i}", level=1)
        doc.add_paragraph(clause)
    doc.add_heading("Signatures", level=1)
    doc.add_paragraph("Acme Kft representative: ____________________")
    doc.add_paragraph("Contractor: ____________________")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Attachment payload registry — keyed by manifest `attachment` field.
# ---------------------------------------------------------------------------

_INVOICE_SPECS = {
    "invoice_pdf_1": (1001, 120_000, "Acme Consulting Kft"),
    "invoice_pdf_2": (1002, 85_000, "BlueSky Solutions Ltd"),
    "invoice_pdf_3": (1003, 45_500, "Overdue Systems Sp. z o.o."),
    "invoice_pdf_4": (1004, 29_900, "Streamflow SaaS Kft"),
    "invoice_pdf_5": (1005, 18_420, "ELMU Nyrt."),
    "invoice_pdf_6": (1006, 210_000, "BlueSky Solutions Ltd"),
    "invoice_pdf_7": (1007, 75_000, "Acme Consulting Kft"),
    "invoice_pdf_8": (1008, 99_000, "Acme Consulting Kft"),
}

_CONTRACT_SPECS = {
    "contract_docx_1": (
        "Service Contract",
        [
            "Contractor shall provide consulting services as described in Schedule A.",
            "Payment terms: net-30 from invoice date, total value 120,000 HUF.",
            "Termination clause: 30 days written notice required by either party.",
        ],
    ),
    "contract_docx_2": (
        "Szolgaltatasi Szerzodes Megujitas 2026",
        [
            "A szerzodo felek kozos megegyezessel meghosszabbitjak a 2025-os szerzodest.",
            "Uj hatalyossag: 2026-04-01 - 2027-03-31.",
            "A szolgaltatasi dij 10%-kal emelkedik az elozo evhez kepest.",
        ],
    ),
    "contract_docx_3": (
        "Non-Disclosure Agreement (NDA)",
        [
            "Confidential Information includes all data marked or treated as confidential.",
            "Term: 3 years from the effective date.",
            "Survival: Confidentiality obligations survive termination for 5 years.",
        ],
    ),
    "contract_docx_4": (
        "Contract Amendment Q2 2026",
        [
            "Amendment #2 to the Master Services Agreement dated 2025-01-15.",
            "Scope: adds a second deployment environment starting 2026-05-01.",
            "Fee adjustment: +15,000 HUF / month for the added environment.",
        ],
    ),
    "contract_docx_5": (
        "Service Level Agreement (SLA)",
        [
            "Uptime commitment: 99.9% measured monthly, excluding planned maintenance.",
            "Response time for P1 incidents: 30 minutes, 24x7.",
            "Penalty clauses: service credit per 0.1% uptime missed.",
        ],
    ),
    "contract_docx_6": (
        "Master Services Agreement (MSA)",
        [
            "This MSA governs all Statements of Work executed between the parties.",
            "Intellectual property: all deliverables assigned to Customer upon payment.",
            "Governing law: Hungary; disputes resolved by the Budapest Arbitration Court.",
        ],
    ),
}

_GENERIC_PDF_SPECS = {
    "other_pdf_1": (
        "Crash Log - Mobile App",
        [
            "Environment: iOS 17.4, app version 2.8.1",
            "Steps to reproduce:",
            "  1. Open upload screen",
            "  2. Select a PDF larger than 5 MB",
            "  3. Tap 'Upload'",
            "Observed: app crashes with SIGSEGV",
            "Expected: upload completes or shows a size-limit error.",
            "Attached this log for the support team reference.",
        ],
    ),
    "other_pdf_2": (
        "Tavaszi katalogus 2026",
        [
            "Uj termekeink a tavaszi szezonra.",
            "30% kedvezmeny minden termekre hetvegen.",
            "Rendeljen online hirlevel-feliratkozoknak.",
            "Kapcsolat: marketing@example.com",
        ],
    ),
}


def _build_attachment(key: str) -> tuple[bytes, str, str, str] | None:
    """Return (payload, maintype, subtype, filename) or None for 'none'."""
    if key == "none":
        return None
    if key.startswith("invoice_pdf_"):
        num, amount, supplier = _INVOICE_SPECS[key]
        payload = _invoice_pdf(num, amount, supplier)
        return payload, "application", "pdf", f"invoice_INV-2026-{num:04d}.pdf"
    if key.startswith("contract_docx_"):
        title, clauses = _CONTRACT_SPECS[key]
        payload = _contract_docx(title, clauses)
        return (
            payload,
            "application",
            "vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{key}.docx",
        )
    if key.startswith("other_pdf_"):
        title, lines = _GENERIC_PDF_SPECS[key]
        payload = _generic_pdf(title, lines)
        return payload, "application", "pdf", f"{key}.pdf"
    raise ValueError(f"Unknown attachment key: {key}")


def _build_eml(entry: dict) -> bytes:
    msg = EmailMessage()
    msg["Subject"] = entry["subject"]
    msg["From"] = "sender@example.com"
    msg["To"] = "intake@example.com"
    msg["Message-ID"] = f"<{entry['id']}@sprint-o.local>"
    msg.set_content(entry["body"])
    att = _build_attachment(entry["attachment"])
    if att is not None:
        payload, maintype, subtype, filename = att
        msg.add_attachment(payload, maintype=maintype, subtype=subtype, filename=filename)
    else:
        # Intake CHECK constraint requires association_mode to be set when
        # intake_descriptions rows exist, which means the package must have at
        # least one file. Mirror the integration-test fixture pattern
        # (`test_scan_and_classify.py::_build_email`) and embed a tiny marker
        # text/plain part so body-only fixtures survive `IntakePackageSink`.
        # The marker is invisible to the classifier (EMAIL_BODY role is
        # extracted only from the main text body) and carries no intent signal.
        msg.add_attachment(
            b"fixture marker\n",
            maintype="text",
            subtype="plain",
            filename="note.txt",
        )
    return msg.as_bytes()


def main() -> int:
    manifest_path = FIXTURE_DIR / "manifest.yaml"
    with manifest_path.open(encoding="utf-8") as f:
        manifest = yaml.safe_load(f)
    fixtures = manifest["fixtures"]
    written = 0
    for entry in fixtures:
        out = FIXTURE_DIR / f"{entry['id']}.eml"
        out.write_bytes(_build_eml(entry))
        written += 1
    print(f"wrote {written} fixture .eml files to {FIXTURE_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
